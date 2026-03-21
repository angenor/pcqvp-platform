import io
import unicodedata

from sqlalchemy.ext.asyncio import AsyncSession

from app.services.account_service import (
    calculate_depenses_recap,
    calculate_equilibre,
    calculate_recettes_recap,
    get_depenses_with_computed,
    get_recettes_with_computed,
)
from app.services.compte_service import get_collectivite_name


def sanitize_filename(name: str) -> str:
    """Replace special characters with ASCII-safe equivalents for filenames."""
    nfkd = unicodedata.normalize("NFKD", name)
    ascii_str = nfkd.encode("ascii", "ignore").decode("ascii")
    safe = ascii_str.replace("'", "").replace(" ", "_")
    return "".join(c for c in safe if c.isalnum() or c in ("_", "-"))


async def generate_excel(
    db: AsyncSession,
    compte,
) -> io.BytesIO:
    """Generate multi-sheet Excel workbook for a compte."""
    from openpyxl import Workbook
    from openpyxl.styles import Border, Font, PatternFill, Side

    wb = Workbook()

    cname = await get_collectivite_name(
        db, compte.collectivite_type, compte.collectivite_id
    )
    bold = Font(bold=True)
    header_fill = PatternFill(
        start_color="D9E2F3", end_color="D9E2F3",
        fill_type="solid",
    )
    section_fill = PatternFill(
        start_color="E2EFDA", end_color="E2EFDA",
        fill_type="solid",
    )
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    num_width = 15

    def _format_taux(val):
        if val is None or val == 0:
            return ""
        return f"{val:.1%}"

    def _section_label(lines, current_idx):
        """Return section header if this line starts a new section."""
        line = lines[current_idx]
        section = line.get("section", "")
        if current_idx == 0:
            return section
        prev_section = lines[current_idx - 1].get("section", "")
        if section != prev_section:
            return section
        return None

    # --- Recettes sheet ---
    recettes = await get_recettes_with_computed(db, compte.id)
    ws_r = wb.active
    ws_r.title = "Recettes"
    ws_r.append([f"Recettes - {cname} - Exercice {compte.annee_exercice}"])
    ws_r.merge_cells("A1:J1")
    ws_r["A1"].font = Font(bold=True, size=14)
    ws_r.append([])

    r_headers = [
        "Compte", "Intitules", "Budget Primitif", "Budget Additionnel",
        "Modifications +/-", "Previsions Definitives", "OR Admis",
        "Recouvrement", "Reste a Recouvrer", "Taux d'Execution",
    ]
    ws_r.append(r_headers)
    for cell in ws_r[ws_r.max_row]:
        cell.font = bold
        cell.fill = header_fill

    r_lines = recettes.get("lines", [])
    for idx, line in enumerate(r_lines):
        sec_label = _section_label(r_lines, idx)
        if sec_label:
            label = (
                "RECETTES DE FONCTIONNEMENT"
                if sec_label == "fonctionnement"
                else "RECETTES D'INVESTISSEMENT"
            )
            ws_r.append([label])
            ws_r.merge_cells(
                start_row=ws_r.max_row, start_column=1,
                end_row=ws_r.max_row, end_column=len(r_headers),
            )
            for cell in ws_r[ws_r.max_row]:
                cell.font = bold
                cell.fill = section_fill

        code = line["compte_code"]
        sums = recettes["hierarchical_sums"].get(code, {})
        vals = sums if sums else line.get("values", {})
        computed = line.get("computed", {})
        indent = "  " * (line["level"] - 1)
        taux = _gv(sums, computed, "taux_execution")
        ws_r.append([
            code,
            f"{indent}{line['intitule']}",
            vals.get("budget_primitif", 0),
            vals.get("budget_additionnel", 0),
            vals.get("modifications", 0),
            _gv(sums, computed, "previsions_definitives"),
            vals.get("or_admis", 0),
            vals.get("recouvrement", 0),
            _gv(sums, computed, "reste_a_recouvrer"),
            _format_taux(taux),
        ])
        if line["level"] == 1:
            for cell in ws_r[ws_r.max_row]:
                cell.font = bold

    # --- Depenses sheets (one per programme) ---
    depenses = await get_depenses_with_computed(db, compte.id)
    for prog in depenses:
        ws_d = wb.create_sheet(title=f"Prog {prog['numero']}")
        title = (
            f"Depenses - Programme {prog['numero']}: "
            f"{prog['intitule']} - {cname} - "
            f"{compte.annee_exercice}"
        )
        ws_d.append([title])
        ws_d.merge_cells("A1:K1")
        ws_d["A1"].font = Font(bold=True, size=14)
        ws_d.append([])

        d_headers = [
            "Compte", "Intitules", "Budget Primitif", "Budget Additionnel",
            "Modifications +/-", "Previsions Definitives", "Engagement",
            "Mandat Admis", "Paiement", "Reste a Payer", "Taux d'Execution",
        ]
        ws_d.append(d_headers)
        for cell in ws_d[ws_d.max_row]:
            cell.font = bold
            cell.fill = header_fill

        d_lines = prog["lines"]
        for idx, line in enumerate(d_lines):
            sec_label = _section_label(d_lines, idx)
            if sec_label:
                label = (
                    "DEPENSES DE FONCTIONNEMENT"
                    if sec_label == "fonctionnement"
                    else "DEPENSES D'INVESTISSEMENT"
                )
                ws_d.append([label])
                ws_d.merge_cells(
                    start_row=ws_d.max_row, start_column=1,
                    end_row=ws_d.max_row, end_column=len(d_headers),
                )
                for cell in ws_d[ws_d.max_row]:
                    cell.font = bold
                    cell.fill = section_fill

            code = line["compte_code"]
            sums = prog["hierarchical_sums"].get(code, {})
            vals = sums if sums else line.get("values", {})
            raw_vals = line.get("values", {})
            computed = line.get("computed", {})
            indent = "  " * (line["level"] - 1)
            taux = _gv(sums, computed, "taux_execution")
            ws_d.append([
                code,
                f"{indent}{line['intitule']}",
                vals.get("budget_primitif", 0),
                vals.get("budget_additionnel", 0),
                vals.get("modifications", 0),
                _gv(sums, computed, "previsions_definitives"),
                raw_vals.get("engagement", 0) or 0,
                vals.get("mandat_admis", 0),
                vals.get("paiement", 0),
                _gv(sums, computed, "reste_a_payer"),
                _format_taux(taux),
            ])
            if line["level"] == 1:
                for cell in ws_d[ws_d.max_row]:
                    cell.font = bold

    # --- Recap Recettes sheet ---
    recap_r = await calculate_recettes_recap(db, compte.id)
    ws_rr = wb.create_sheet(title="Recap Recettes")
    ws_rr.append(["Recapitulatif des recettes"])
    ws_rr["A1"].font = Font(bold=True, size=14)
    ws_rr.append([])
    recap_r_headers = [
        "Compte", "Intitules", "Previsions Definitives",
        "OR Admis", "Recouvrement", "Reste a Recouvrer",
    ]
    ws_rr.append(recap_r_headers)
    for cell in ws_rr[ws_rr.max_row]:
        cell.font = bold
        cell.fill = header_fill

    for section in recap_r.get("sections", []):
        sec_label = (
            "RECETTES DE FONCTIONNEMENT"
            if section["section"] == "fonctionnement"
            else "RECETTES D'INVESTISSEMENT"
        )
        ws_rr.append([sec_label])
        for cell in ws_rr[ws_rr.max_row]:
            cell.font = bold
            cell.fill = section_fill

        for cat in section.get("categories", []):
            ws_rr.append([
                cat["compte_code"],
                cat["intitule"],
                cat.get("previsions_definitives", 0),
                cat.get("or_admis", 0),
                cat.get("recouvrement", 0),
                cat.get("reste_a_recouvrer", 0),
            ])

        total = section.get("total_section", {})
        ws_rr.append([
            "",
            f"Total {section['section']}",
            total.get("previsions_definitives", 0),
            total.get("or_admis", 0),
            total.get("recouvrement", 0),
            total.get("reste_a_recouvrer", 0),
        ])
        for cell in ws_rr[ws_rr.max_row]:
            cell.font = bold

    # --- Recap Depenses par Programme sheet ---
    recap_d = await calculate_depenses_recap(db, compte.id)
    progs = recap_d.get("programmes", [])

    ws_rdp = wb.create_sheet(title="Recap Dep par Prog")
    ws_rdp.append(["Recapitulatif des depenses par programme"])
    ws_rdp["A1"].font = Font(bold=True, size=14)
    ws_rdp.append([])

    prog_labels = [f"Prog {p.get('numero', '')}" for p in progs]
    rdp_headers = ["Compte", "Intitules"]
    for pl in prog_labels:
        rdp_headers.append(f"Mand. {pl}")
    rdp_headers.append("Total Mand.")
    for pl in prog_labels:
        rdp_headers.append(f"Paie. {pl}")
    rdp_headers.append("Total Paie.")
    for pl in prog_labels:
        rdp_headers.append(f"Reste {pl}")
    rdp_headers.append("Total Reste")

    ws_rdp.append(rdp_headers)
    for cell in ws_rdp[ws_rdp.max_row]:
        cell.font = bold
        cell.fill = header_fill

    for section in recap_d.get("sections", []):
        sec_label = (
            "DEPENSES DE FONCTIONNEMENT"
            if section["section"] == "fonctionnement"
            else "DEPENSES D'INVESTISSEMENT"
        )
        ws_rdp.append([sec_label])
        ws_rdp.merge_cells(
            start_row=ws_rdp.max_row, start_column=1,
            end_row=ws_rdp.max_row, end_column=len(rdp_headers),
        )
        for cell in ws_rdp[ws_rdp.max_row]:
            cell.font = bold
            cell.fill = section_fill

        for cat in section.get("categories", []):
            row = [cat["compte_code"], cat["intitule"]]
            cat_progs = cat.get("programmes", [])
            cat_total = cat.get("total", {})
            for p in cat_progs:
                row.append(p.get("mandat_admis", 0))
            row.append(cat_total.get("mandat_admis", 0))
            for p in cat_progs:
                row.append(p.get("paiement", 0))
            row.append(cat_total.get("paiement", 0))
            for p in cat_progs:
                row.append(p.get("reste_a_payer", 0))
            row.append(cat_total.get("reste_a_payer", 0))
            ws_rdp.append(row)

        total_s = section.get("total_section", {})
        row = ["", f"Total {section['section']}"]
        row += [""] * len(progs)
        row.append(total_s.get("mandat_admis", 0))
        row += [""] * len(progs)
        row.append(total_s.get("paiement", 0))
        row += [""] * len(progs)
        row.append(total_s.get("reste_a_payer", 0))
        ws_rdp.append(row)
        for cell in ws_rdp[ws_rdp.max_row]:
            cell.font = bold

    # --- Recap Depenses sheet (synthese) ---
    ws_rd = wb.create_sheet(title="Recap Depenses")
    ws_rd.append(["Recapitulatif des depenses"])
    ws_rd["A1"].font = Font(bold=True, size=14)
    ws_rd.append([])

    rd_headers = [
        "Compte", "Intitules", "Mandat Admis", "Paiement", "Reste a Payer",
    ]
    ws_rd.append(rd_headers)
    for cell in ws_rd[ws_rd.max_row]:
        cell.font = bold
        cell.fill = header_fill

    for section in recap_d.get("sections", []):
        sec_label = (
            "DEPENSES DE FONCTIONNEMENT"
            if section["section"] == "fonctionnement"
            else "DEPENSES D'INVESTISSEMENT"
        )
        ws_rd.append([sec_label])
        for cell in ws_rd[ws_rd.max_row]:
            cell.font = bold
            cell.fill = section_fill

        for cat in section.get("categories", []):
            cat_total = cat.get("total", {})
            ws_rd.append([
                cat["compte_code"],
                cat["intitule"],
                cat_total.get("mandat_admis", 0),
                cat_total.get("paiement", 0),
                cat_total.get("reste_a_payer", 0),
            ])

        total_s = section.get("total_section", {})
        ws_rd.append([
            "",
            f"Total {section['section']}",
            total_s.get("mandat_admis", 0),
            total_s.get("paiement", 0),
            total_s.get("reste_a_payer", 0),
        ])
        for cell in ws_rd[ws_rd.max_row]:
            cell.font = bold

    # --- Equilibre sheet ---
    eq = await calculate_equilibre(db, compte.id)
    ws_eq = wb.create_sheet(title="Equilibre")
    ws_eq.append(["Equilibre budgetaire"])
    ws_eq.merge_cells("A1:J1")
    ws_eq["A1"].font = Font(bold=True, size=14)
    ws_eq.append([])

    eq_headers = [
        "Compte", "Intitules", "Mandat Admis", "Paiement", "Reste a Payer",
        "Compte", "Intitules", "OR Admis", "Recouvrement", "Reste a Recouvrer",
    ]
    ws_eq.append(eq_headers)
    for cell in ws_eq[ws_eq.max_row]:
        cell.font = bold
        cell.fill = header_fill

    for sec_name in ("fonctionnement", "investissement"):
        sec = eq.get(sec_name, {})
        sec_label = sec_name.upper()
        ws_eq.append([sec_label])
        ws_eq.merge_cells(
            start_row=ws_eq.max_row, start_column=1,
            end_row=ws_eq.max_row, end_column=len(eq_headers),
        )
        for cell in ws_eq[ws_eq.max_row]:
            cell.font = bold
            cell.fill = section_fill

        dep = sec.get("depenses", {})
        rec = sec.get("recettes", {})
        dep_r = dep.get("reelles", [])
        rec_r = rec.get("reelles", [])
        max_len = max(len(dep_r), len(rec_r), 1)

        for i in range(max_len):
            d_item = dep_r[i] if i < len(dep_r) else {}
            r_item = rec_r[i] if i < len(rec_r) else {}
            ws_eq.append([
                d_item.get("compte_code", ""),
                d_item.get("intitule", ""),
                d_item.get("montant", ""),
                "", "",
                r_item.get("compte_code", ""),
                r_item.get("intitule", ""),
                r_item.get("montant", ""),
                "", "",
            ])

        # Sous-total
        ws_eq.append([
            "", f"Total {sec_name}",
            dep.get("total", 0), "", "",
            "", f"Total {sec_name}",
            rec.get("total", 0), "", "",
        ])
        for cell in ws_eq[ws_eq.max_row]:
            cell.font = bold

        # Operations d'ordre
        dep_o = dep.get("ordre", [])
        rec_o = rec.get("ordre", [])
        if dep_o or rec_o:
            max_o = max(len(dep_o), len(rec_o), 1)
            ws_eq.append(["", "Operations d'ordre", "", "", "",
                          "", "Operations d'ordre", "", "", ""])
            for cell in ws_eq[ws_eq.max_row]:
                cell.font = bold
            for i in range(max_o):
                do = dep_o[i] if i < len(dep_o) else {}
                ro = rec_o[i] if i < len(rec_o) else {}
                ws_eq.append([
                    do.get("compte_code", ""),
                    do.get("intitule", ""),
                    do.get("montant", ""),
                    "", "",
                    ro.get("compte_code", ""),
                    ro.get("intitule", ""),
                    ro.get("montant", ""),
                    "", "",
                ])

        excedent = sec.get("excedent", 0)
        label = "Excedent" if excedent >= 0 else "Deficit"
        ws_eq.append(["", label, abs(excedent), "", "",
                       "", "", "", "", ""])
        for cell in ws_eq[ws_eq.max_row]:
            cell.font = bold
        ws_eq.append([])

    rd = eq.get("resultat_definitif", 0)
    rd_label = "EXCEDENT DEFINITIF" if rd >= 0 else "DEFICIT DEFINITIF"
    ws_eq.append([rd_label, "", abs(rd)])
    for cell in ws_eq[ws_eq.max_row]:
        cell.font = Font(bold=True, size=12)

    # --- Apply formatting to all sheets ---
    for ws in wb.worksheets:
        _apply_formatting(ws, thin_border, num_width)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def _gv(sums: dict, computed: dict, key: str):
    """Get value from sums first, then computed."""
    return sums.get(key, computed.get(key, 0))


def _apply_formatting(ws, thin_border, num_width):
    """Apply borders and column widths to a worksheet."""
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row,
                            max_col=ws.max_column):
        for cell in row:
            if cell.row >= 3 and cell.value is not None:
                cell.border = thin_border

    for col_cells in ws.columns:
        col_letter = col_cells[0].column_letter
        col_idx = col_cells[0].column
        if col_idx == 1:
            ws.column_dimensions[col_letter].width = 12
        elif col_idx == 2:
            max_len = 0
            for cell in col_cells:
                val = str(cell.value or "")
                max_len = max(max_len, len(val))
            ws.column_dimensions[col_letter].width = min(max_len + 2, 50)
        else:
            ws.column_dimensions[col_letter].width = num_width


async def generate_word(
    db: AsyncSession,
    compte,
) -> io.BytesIO:
    """Generate a Word document for a compte."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    cname = await get_collectivite_name(
        db, compte.collectivite_type, compte.collectivite_id
    )

    doc.add_heading(
        f"Compte administratif - {cname} - Exercice {compte.annee_exercice}",
        level=0,
    )
    doc.add_paragraph(
        f"Type: {compte.collectivite_type.capitalize()} | "
        f"Statut: {compte.status.value}"
    )

    def _format_taux_word(val):
        if val is None or val == 0:
            return ""
        return f"{val:.1%}"

    def _set_bold_row(row):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    def _add_row_data(table, cells_data, is_bold=False, font_size=None):
        row = table.add_row()
        for i, val in enumerate(cells_data):
            row.cells[i].text = str(val)
        if is_bold:
            _set_bold_row(row)
        if font_size:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(font_size)
        return row

    # --- Recettes ---
    doc.add_heading("Recettes", level=1)
    recettes = await get_recettes_with_computed(db, compte.id)

    r_headers = [
        "Compte", "Intitules", "Budget Primitif", "Budget Additionnel",
        "Modifications +/-", "Prev. Def.", "OR Admis",
        "Recouvrement", "Reste a Recouvrer", "Taux Exec.",
    ]
    table_r = doc.add_table(rows=1, cols=len(r_headers))
    table_r.style = "Table Grid"
    for i, h in enumerate(r_headers):
        table_r.rows[0].cells[i].text = h
    _set_bold_row(table_r.rows[0])

    r_lines = recettes.get("lines", [])
    current_section = None
    for line in r_lines:
        section = line.get("section", "")
        if section != current_section:
            current_section = section
            label = (
                "RECETTES DE FONCTIONNEMENT"
                if section == "fonctionnement"
                else "RECETTES D'INVESTISSEMENT"
            )
            _add_row_data(table_r, [label] + [""] * (len(r_headers) - 1),
                          is_bold=True)

        code = line["compte_code"]
        sums = recettes["hierarchical_sums"].get(code, {})
        vals = sums if sums else line.get("values", {})
        computed = line.get("computed", {})
        indent = "  " * (line["level"] - 1)
        taux = _gv(sums, computed, "taux_execution")

        row_data = [
            code,
            f"{indent}{line['intitule']}",
            vals.get("budget_primitif", 0),
            vals.get("budget_additionnel", 0),
            vals.get("modifications", 0),
            _gv(sums, computed, "previsions_definitives"),
            vals.get("or_admis", 0),
            vals.get("recouvrement", 0),
            _gv(sums, computed, "reste_a_recouvrer"),
            _format_taux_word(taux),
        ]
        _add_row_data(table_r, row_data, is_bold=(line["level"] == 1))

    # --- Depenses ---
    depenses = await get_depenses_with_computed(db, compte.id)
    for prog in depenses:
        doc.add_heading(
            f"Depenses - Programme {prog['numero']}: {prog['intitule']}",
            level=1,
        )
        d_headers = [
            "Compte", "Intitules", "Budget Primitif", "Budget Additionnel",
            "Modifications +/-", "Prev. Def.", "Engagement",
            "Mandat Admis", "Paiement", "Reste a Payer", "Taux Exec.",
        ]
        table_d = doc.add_table(rows=1, cols=len(d_headers))
        table_d.style = "Table Grid"
        for i, h in enumerate(d_headers):
            table_d.rows[0].cells[i].text = h
        _set_bold_row(table_d.rows[0])

        d_lines = prog["lines"]
        current_section = None
        for line in d_lines:
            section = line.get("section", "")
            if section != current_section:
                current_section = section
                label = (
                    "DEPENSES DE FONCTIONNEMENT"
                    if section == "fonctionnement"
                    else "DEPENSES D'INVESTISSEMENT"
                )
                _add_row_data(table_d,
                              [label] + [""] * (len(d_headers) - 1),
                              is_bold=True)

            code = line["compte_code"]
            sums = prog["hierarchical_sums"].get(code, {})
            vals = sums if sums else line.get("values", {})
            raw_vals = line.get("values", {})
            computed = line.get("computed", {})
            indent = "  " * (line["level"] - 1)
            taux = _gv(sums, computed, "taux_execution")

            row_data = [
                code,
                f"{indent}{line['intitule']}",
                vals.get("budget_primitif", 0),
                vals.get("budget_additionnel", 0),
                vals.get("modifications", 0),
                _gv(sums, computed, "previsions_definitives"),
                raw_vals.get("engagement", 0) or 0,
                vals.get("mandat_admis", 0),
                vals.get("paiement", 0),
                _gv(sums, computed, "reste_a_payer"),
                _format_taux_word(taux),
            ]
            _add_row_data(table_d, row_data, is_bold=(line["level"] == 1))

    # --- Recapitulatif Recettes ---
    recap_r = await calculate_recettes_recap(db, compte.id)
    doc.add_heading("Recapitulatif des recettes", level=1)

    rr_headers = [
        "Compte", "Intitules", "Previsions Definitives",
        "OR Admis", "Recouvrement", "Reste a Recouvrer",
    ]
    table_rr = doc.add_table(rows=1, cols=len(rr_headers))
    table_rr.style = "Table Grid"
    for i, h in enumerate(rr_headers):
        table_rr.rows[0].cells[i].text = h
    _set_bold_row(table_rr.rows[0])

    for section in recap_r.get("sections", []):
        sec_label = (
            "RECETTES DE FONCTIONNEMENT"
            if section["section"] == "fonctionnement"
            else "RECETTES D'INVESTISSEMENT"
        )
        _add_row_data(table_rr,
                      [sec_label] + [""] * (len(rr_headers) - 1),
                      is_bold=True)

        for cat in section.get("categories", []):
            _add_row_data(table_rr, [
                cat["compte_code"],
                cat["intitule"],
                cat.get("previsions_definitives", 0),
                cat.get("or_admis", 0),
                cat.get("recouvrement", 0),
                cat.get("reste_a_recouvrer", 0),
            ])

        total = section.get("total_section", {})
        _add_row_data(table_rr, [
            "",
            f"Total {section['section']}",
            total.get("previsions_definitives", 0),
            total.get("or_admis", 0),
            total.get("recouvrement", 0),
            total.get("reste_a_recouvrer", 0),
        ], is_bold=True)

    # --- Recapitulatif Depenses ---
    recap_d = await calculate_depenses_recap(db, compte.id)
    doc.add_heading("Recapitulatif des depenses", level=1)

    rd_headers = [
        "Compte", "Intitules", "Mandat Admis", "Paiement", "Reste a Payer",
    ]
    table_rd = doc.add_table(rows=1, cols=len(rd_headers))
    table_rd.style = "Table Grid"
    for i, h in enumerate(rd_headers):
        table_rd.rows[0].cells[i].text = h
    _set_bold_row(table_rd.rows[0])

    for section in recap_d.get("sections", []):
        sec_label = (
            "DEPENSES DE FONCTIONNEMENT"
            if section["section"] == "fonctionnement"
            else "DEPENSES D'INVESTISSEMENT"
        )
        _add_row_data(table_rd,
                      [sec_label] + [""] * (len(rd_headers) - 1),
                      is_bold=True)

        for cat in section.get("categories", []):
            cat_total = cat.get("total", {})
            _add_row_data(table_rd, [
                cat["compte_code"],
                cat["intitule"],
                cat_total.get("mandat_admis", 0),
                cat_total.get("paiement", 0),
                cat_total.get("reste_a_payer", 0),
            ])

        total_s = section.get("total_section", {})
        _add_row_data(table_rd, [
            "",
            f"Total {section['section']}",
            total_s.get("mandat_admis", 0),
            total_s.get("paiement", 0),
            total_s.get("reste_a_payer", 0),
        ], is_bold=True)

    # --- Equilibre ---
    eq = await calculate_equilibre(db, compte.id)
    doc.add_heading("Equilibre budgetaire", level=1)

    for sec_name in ("fonctionnement", "investissement"):
        sec = eq.get(sec_name, {})
        doc.add_heading(sec_name.capitalize(), level=2)

        table_eq = doc.add_table(rows=1, cols=3)
        table_eq.style = "Table Grid"
        table_eq.rows[0].cells[0].text = ""
        table_eq.rows[0].cells[1].text = "Depenses"
        table_eq.rows[0].cells[2].text = "Recettes"
        _set_bold_row(table_eq.rows[0])

        dep = sec.get("depenses", {})
        rec = sec.get("recettes", {})
        dep_r = dep.get("reelles", [])
        rec_r = rec.get("reelles", [])
        max_len = max(len(dep_r), len(rec_r), 1)

        for i in range(max_len):
            d_item = dep_r[i] if i < len(dep_r) else {}
            r_item = rec_r[i] if i < len(rec_r) else {}
            label = (
                d_item.get("intitule", "")
                or r_item.get("intitule", "")
            )
            _add_row_data(table_eq, [
                label,
                d_item.get("montant", ""),
                r_item.get("montant", ""),
            ])

        _add_row_data(table_eq, [
            "Total", dep.get("total", 0), rec.get("total", 0),
        ], is_bold=True)

        excedent = sec.get("excedent", 0)
        label = "Excedent" if excedent >= 0 else "Deficit"
        doc.add_paragraph(f"{label}: {abs(excedent):,.0f}")

    doc.add_heading("Resultat definitif", level=1)
    rd = eq.get("resultat_definitif", 0)
    doc.add_paragraph(f"{'Excedent' if rd >= 0 else 'Deficit'}: {abs(rd):,.0f}")

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
